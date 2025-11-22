"""
Custom Report Dialog
Allows users to configure custom report sections, logo, and export options.
"""

from PyQt5.QtWidgets import (
    QDialog, QVBoxLayout, QHBoxLayout, QLabel, QLineEdit,
    QTextEdit, QPushButton, QFileDialog, QGroupBox, QFormLayout,
    QMessageBox, QTabWidget, QWidget, QComboBox, QCheckBox
)
from PyQt5.QtCore import Qt
from PyQt5.QtGui import QPixmap
from pathlib import Path
import json
import os
from datetime import datetime


class CustomReportDialog(QDialog):
    """Dialog for configuring custom report content and exporting"""

    def __init__(self, results_data=None, parent=None):
        super().__init__(parent)
        self.results_data = results_data or {}
        self.logo_path = None
        self.setWindowTitle("Custom Report Configuration")
        self.resize(700, 600)
        self.init_ui()
        self.load_saved_config()

    def init_ui(self):
        """Initialize the dialog UI"""
        self.setStyleSheet("""
            QDialog {
                background-color: #ffffff;
                color: #333333;
            }
            QLabel {
                color: #333333;
            }
            QGroupBox {
                background-color: #f8f8f8;
                border: 1px solid #e0e0e0;
                border-radius: 6px;
                margin-top: 10px;
                padding-top: 10px;
                font-weight: bold;
                color: #333333;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 10px;
                padding: 0 5px;
                color: #4CAF50;
            }
            QLineEdit, QTextEdit {
                background-color: #ffffff;
                color: #333333;
                border: 1px solid #cccccc;
                border-radius: 4px;
                padding: 6px;
            }
            QLineEdit:focus, QTextEdit:focus {
                border: 2px solid #4CAF50;
            }
            QPushButton {
                background-color: #4CAF50;
                color: white;
                border: none;
                padding: 8px 16px;
                border-radius: 4px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #45a049;
            }
        """)

        layout = QVBoxLayout(self)

        # Create tabs for different sections
        tabs = QTabWidget()
        tabs.setStyleSheet("""
            QTabWidget::pane {
                border: 1px solid #e0e0e0;
                background-color: #ffffff;
            }
            QTabBar::tab {
                background-color: #f5f5f5;
                color: #333333;
                padding: 8px 15px;
                border: 1px solid #e0e0e0;
                border-bottom: none;
            }
            QTabBar::tab:selected {
                background-color: #ffffff;
                color: #4CAF50;
                font-weight: bold;
            }
        """)

        # General tab
        general_tab = self._create_general_tab()
        tabs.addTab(general_tab, "General")

        # Content tab
        content_tab = self._create_content_tab()
        tabs.addTab(content_tab, "Content Sections")

        # Export tab
        export_tab = self._create_export_tab()
        tabs.addTab(export_tab, "Export Options")

        layout.addWidget(tabs)

        # Buttons
        btn_layout = QHBoxLayout()
        btn_layout.addStretch()

        save_config_btn = QPushButton("Save Configuration")
        save_config_btn.setStyleSheet("""
            QPushButton {
                background-color: #2196F3;
                color: white;
                border: none;
                padding: 8px 16px;
                border-radius: 4px;
            }
            QPushButton:hover {
                background-color: #1976D2;
            }
        """)
        save_config_btn.clicked.connect(self.save_config)
        btn_layout.addWidget(save_config_btn)

        generate_btn = QPushButton("Generate Report")
        generate_btn.clicked.connect(self.generate_report)
        btn_layout.addWidget(generate_btn)

        cancel_btn = QPushButton("Cancel")
        cancel_btn.setStyleSheet("""
            QPushButton {
                background-color: #f5f5f5;
                color: #333333;
                border: 1px solid #cccccc;
            }
            QPushButton:hover {
                background-color: #e0e0e0;
            }
        """)
        cancel_btn.clicked.connect(self.reject)
        btn_layout.addWidget(cancel_btn)

        layout.addLayout(btn_layout)

    def _create_general_tab(self):
        """Create general settings tab"""
        widget = QWidget()
        layout = QVBoxLayout(widget)

        # Logo section
        logo_group = QGroupBox("Company Logo")
        logo_layout = QVBoxLayout()

        logo_row = QHBoxLayout()
        self.logo_preview = QLabel("No logo selected")
        self.logo_preview.setAlignment(Qt.AlignCenter)
        self.logo_preview.setMinimumHeight(100)
        self.logo_preview.setStyleSheet("border: 1px dashed #cccccc; border-radius: 4px;")
        logo_row.addWidget(self.logo_preview)

        logo_btn_layout = QVBoxLayout()
        select_logo_btn = QPushButton("Select Logo")
        select_logo_btn.clicked.connect(self.select_logo)
        logo_btn_layout.addWidget(select_logo_btn)

        clear_logo_btn = QPushButton("Clear")
        clear_logo_btn.setStyleSheet("""
            QPushButton {
                background-color: #f44336;
            }
            QPushButton:hover {
                background-color: #d32f2f;
            }
        """)
        clear_logo_btn.clicked.connect(self.clear_logo)
        logo_btn_layout.addWidget(clear_logo_btn)

        logo_btn_layout.addStretch()
        logo_row.addLayout(logo_btn_layout)
        logo_layout.addLayout(logo_row)

        logo_group.setLayout(logo_layout)
        layout.addWidget(logo_group)

        # Report info
        info_group = QGroupBox("Report Information")
        info_layout = QFormLayout()

        self.report_title = QLineEdit()
        self.report_title.setPlaceholderText("Vulnerability Assessment Report")
        info_layout.addRow("Report Title:", self.report_title)

        self.company_name = QLineEdit()
        self.company_name.setPlaceholderText("Your Company Name")
        info_layout.addRow("Company Name:", self.company_name)

        self.author_name = QLineEdit()
        self.author_name.setPlaceholderText("Auditor Name")
        info_layout.addRow("Author:", self.author_name)

        self.client_name = QLineEdit()
        self.client_name.setPlaceholderText("Client Name")
        info_layout.addRow("Client:", self.client_name)

        info_group.setLayout(info_layout)
        layout.addWidget(info_group)

        layout.addStretch()
        return widget

    def _create_content_tab(self):
        """Create content sections tab"""
        widget = QWidget()
        layout = QVBoxLayout(widget)

        # Intro section
        intro_group = QGroupBox("Introduction")
        intro_layout = QVBoxLayout()
        self.intro_text = QTextEdit()
        self.intro_text.setPlaceholderText("Enter introduction text for the report...")
        self.intro_text.setMaximumHeight(80)
        intro_layout.addWidget(self.intro_text)
        intro_group.setLayout(intro_layout)
        layout.addWidget(intro_group)

        # Scope section
        scope_group = QGroupBox("Scope")
        scope_layout = QVBoxLayout()
        self.scope_text = QTextEdit()
        self.scope_text.setPlaceholderText("Define the scope of the assessment...")
        self.scope_text.setMaximumHeight(80)
        scope_layout.addWidget(self.scope_text)
        scope_group.setLayout(scope_layout)
        layout.addWidget(scope_group)

        # Team section
        team_group = QGroupBox("Team")
        team_layout = QVBoxLayout()
        self.team_text = QTextEdit()
        self.team_text.setPlaceholderText("List team members involved...")
        self.team_text.setMaximumHeight(60)
        team_layout.addWidget(self.team_text)
        team_group.setLayout(team_layout)
        layout.addWidget(team_group)

        # Summary section
        summary_group = QGroupBox("Executive Summary")
        summary_layout = QVBoxLayout()
        self.summary_text = QTextEdit()
        self.summary_text.setPlaceholderText("Executive summary of findings...")
        self.summary_text.setMaximumHeight(80)
        summary_layout.addWidget(self.summary_text)
        summary_group.setLayout(summary_layout)
        layout.addWidget(summary_group)

        # Methodology section
        methodology_group = QGroupBox("Methodology")
        methodology_layout = QVBoxLayout()
        self.methodology_text = QTextEdit()
        self.methodology_text.setPlaceholderText("Describe the testing methodology...")
        self.methodology_text.setMaximumHeight(80)
        methodology_layout.addWidget(self.methodology_text)
        methodology_group.setLayout(methodology_layout)
        layout.addWidget(methodology_group)

        # Process section
        process_group = QGroupBox("Process")
        process_layout = QVBoxLayout()
        self.process_text = QTextEdit()
        self.process_text.setPlaceholderText("Describe the testing process...")
        self.process_text.setMaximumHeight(80)
        process_layout.addWidget(self.process_text)
        process_group.setLayout(process_layout)
        layout.addWidget(process_group)

        return widget

    def _create_export_tab(self):
        """Create export options tab"""
        widget = QWidget()
        layout = QVBoxLayout(widget)

        # Export format
        format_group = QGroupBox("Export Format")
        format_layout = QVBoxLayout()

        self.export_html = QCheckBox("HTML Report")
        self.export_html.setChecked(True)
        format_layout.addWidget(self.export_html)

        self.export_pdf = QCheckBox("PDF Report")
        format_layout.addWidget(self.export_pdf)

        self.export_docx = QCheckBox("DOCX Report (Word)")
        format_layout.addWidget(self.export_docx)

        self.export_csv = QCheckBox("CSV (Findings Only)")
        format_layout.addWidget(self.export_csv)

        format_group.setLayout(format_layout)
        layout.addWidget(format_group)

        # Include options
        include_group = QGroupBox("Include in Report")
        include_layout = QVBoxLayout()

        self.include_evidence = QCheckBox("Include Evidence/Screenshots")
        self.include_evidence.setChecked(True)
        include_layout.addWidget(self.include_evidence)

        self.include_remediation = QCheckBox("Include Remediation Steps")
        self.include_remediation.setChecked(True)
        include_layout.addWidget(self.include_remediation)

        self.include_references = QCheckBox("Include References (CWE, OWASP)")
        self.include_references.setChecked(True)
        include_layout.addWidget(self.include_references)

        self.include_timeline = QCheckBox("Include Timeline Chart")
        self.include_timeline.setChecked(True)
        include_layout.addWidget(self.include_timeline)

        include_group.setLayout(include_layout)
        layout.addWidget(include_group)

        layout.addStretch()
        return widget

    def select_logo(self):
        """Open file dialog to select logo"""
        filename, _ = QFileDialog.getOpenFileName(
            self, "Select Logo",
            "", "Image Files (*.png *.jpg *.jpeg *.gif *.bmp)"
        )
        if filename:
            self.logo_path = filename
            pixmap = QPixmap(filename)
            if not pixmap.isNull():
                scaled = pixmap.scaled(200, 100, Qt.KeepAspectRatio, Qt.SmoothTransformation)
                self.logo_preview.setPixmap(scaled)
            else:
                self.logo_preview.setText(Path(filename).name)

    def clear_logo(self):
        """Clear the selected logo"""
        self.logo_path = None
        self.logo_preview.setPixmap(QPixmap())
        self.logo_preview.setText("No logo selected")

    def get_config(self):
        """Get the current configuration"""
        return {
            'logo_path': self.logo_path,
            'report_title': self.report_title.text(),
            'company_name': self.company_name.text(),
            'author_name': self.author_name.text(),
            'client_name': self.client_name.text(),
            'intro': self.intro_text.toPlainText(),
            'scope': self.scope_text.toPlainText(),
            'team': self.team_text.toPlainText(),
            'summary': self.summary_text.toPlainText(),
            'methodology': self.methodology_text.toPlainText(),
            'process': self.process_text.toPlainText(),
            'export_html': self.export_html.isChecked(),
            'export_pdf': self.export_pdf.isChecked(),
            'export_docx': self.export_docx.isChecked(),
            'export_csv': self.export_csv.isChecked(),
            'include_evidence': self.include_evidence.isChecked(),
            'include_remediation': self.include_remediation.isChecked(),
            'include_references': self.include_references.isChecked(),
            'include_timeline': self.include_timeline.isChecked()
        }

    def save_config(self):
        """Save configuration to file"""
        config = self.get_config()
        config_path = Path(__file__).parent.parent.parent / "report_config.json"

        try:
            with open(config_path, 'w', encoding='utf-8') as f:
                json.dump(config, f, indent=2)
            QMessageBox.information(self, "Saved", "Report configuration saved successfully.")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to save configuration: {e}")

    def load_saved_config(self):
        """Load saved configuration"""
        config_path = Path(__file__).parent.parent.parent / "report_config.json"

        if config_path.exists():
            try:
                with open(config_path, 'r', encoding='utf-8') as f:
                    config = json.load(f)

                if config.get('logo_path') and os.path.exists(config['logo_path']):
                    self.logo_path = config['logo_path']
                    pixmap = QPixmap(self.logo_path)
                    if not pixmap.isNull():
                        scaled = pixmap.scaled(200, 100, Qt.KeepAspectRatio, Qt.SmoothTransformation)
                        self.logo_preview.setPixmap(scaled)

                self.report_title.setText(config.get('report_title', ''))
                self.company_name.setText(config.get('company_name', ''))
                self.author_name.setText(config.get('author_name', ''))
                self.client_name.setText(config.get('client_name', ''))
                self.intro_text.setPlainText(config.get('intro', ''))
                self.scope_text.setPlainText(config.get('scope', ''))
                self.team_text.setPlainText(config.get('team', ''))
                self.summary_text.setPlainText(config.get('summary', ''))
                self.methodology_text.setPlainText(config.get('methodology', ''))
                self.process_text.setPlainText(config.get('process', ''))
                self.export_html.setChecked(config.get('export_html', True))
                self.export_pdf.setChecked(config.get('export_pdf', False))
                self.export_docx.setChecked(config.get('export_docx', False))
                self.export_csv.setChecked(config.get('export_csv', False))
                self.include_evidence.setChecked(config.get('include_evidence', True))
                self.include_remediation.setChecked(config.get('include_remediation', True))
                self.include_references.setChecked(config.get('include_references', True))
                self.include_timeline.setChecked(config.get('include_timeline', True))
            except:
                pass

    def generate_report(self):
        """Generate the report with selected options"""
        config = self.get_config()

        # At least one format must be selected
        if not any([config['export_html'], config['export_pdf'],
                   config['export_docx'], config['export_csv']]):
            QMessageBox.warning(self, "No Format", "Please select at least one export format.")
            return

        # Get save location
        default_name = f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        save_dir = QFileDialog.getExistingDirectory(
            self, "Select Output Directory",
            str(Path.home())
        )

        if not save_dir:
            return

        try:
            generated = []

            if config['export_html']:
                html_path = Path(save_dir) / f"{default_name}.html"
                self._generate_html(html_path, config)
                generated.append(str(html_path))

            if config['export_csv']:
                csv_path = Path(save_dir) / f"{default_name}.csv"
                self._generate_csv(csv_path)
                generated.append(str(csv_path))

            if config['export_pdf']:
                pdf_path = Path(save_dir) / f"{default_name}.pdf"
                success = self._generate_pdf(pdf_path, config)
                if success:
                    generated.append(str(pdf_path))

            if config['export_docx']:
                docx_path = Path(save_dir) / f"{default_name}.docx"
                success = self._generate_docx(docx_path, config)
                if success:
                    generated.append(str(docx_path))

            if generated:
                QMessageBox.information(
                    self, "Report Generated",
                    f"Generated {len(generated)} report(s):\n\n" + "\n".join(generated)
                )
                self.accept()

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to generate report: {e}")

    def _generate_html(self, path, config):
        """Generate HTML report"""
        html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{config.get('report_title', 'Security Assessment Report')}</title>
    <style>
        body {{ font-family: Arial, sans-serif; margin: 40px; color: #333; }}
        .header {{ text-align: center; margin-bottom: 40px; }}
        .logo {{ max-width: 200px; margin-bottom: 20px; }}
        h1 {{ color: #4CAF50; }}
        h2 {{ color: #2196F3; border-bottom: 2px solid #e0e0e0; padding-bottom: 10px; }}
        .section {{ margin-bottom: 30px; }}
        .meta {{ color: #666; font-size: 14px; margin-bottom: 20px; }}
        table {{ width: 100%; border-collapse: collapse; margin: 20px 0; }}
        th, td {{ border: 1px solid #e0e0e0; padding: 10px; text-align: left; }}
        th {{ background-color: #f5f5f5; }}
        .critical {{ color: #f44336; font-weight: bold; }}
        .high {{ color: #FF9800; font-weight: bold; }}
        .medium {{ color: #FFC107; }}
        .low {{ color: #4CAF50; }}
    </style>
</head>
<body>
    <div class="header">
        {'<img class="logo" src="' + config['logo_path'] + '" alt="Logo">' if config.get('logo_path') else ''}
        <h1>{config.get('report_title', 'Security Assessment Report')}</h1>
        <div class="meta">
            <p><strong>Company:</strong> {config.get('company_name', 'N/A')}</p>
            <p><strong>Client:</strong> {config.get('client_name', 'N/A')}</p>
            <p><strong>Author:</strong> {config.get('author_name', 'N/A')}</p>
            <p><strong>Date:</strong> {datetime.now().strftime('%Y-%m-%d')}</p>
        </div>
    </div>
"""

        # Add custom sections if provided
        if config.get('intro'):
            html += f"""
    <div class="section">
        <h2>Introduction</h2>
        <p>{config['intro'].replace(chr(10), '<br>')}</p>
    </div>
"""

        if config.get('scope'):
            html += f"""
    <div class="section">
        <h2>Scope</h2>
        <p>{config['scope'].replace(chr(10), '<br>')}</p>
    </div>
"""

        if config.get('team'):
            html += f"""
    <div class="section">
        <h2>Team</h2>
        <p>{config['team'].replace(chr(10), '<br>')}</p>
    </div>
"""

        if config.get('summary'):
            html += f"""
    <div class="section">
        <h2>Executive Summary</h2>
        <p>{config['summary'].replace(chr(10), '<br>')}</p>
    </div>
"""

        if config.get('methodology'):
            html += f"""
    <div class="section">
        <h2>Methodology</h2>
        <p>{config['methodology'].replace(chr(10), '<br>')}</p>
    </div>
"""

        if config.get('process'):
            html += f"""
    <div class="section">
        <h2>Process</h2>
        <p>{config['process'].replace(chr(10), '<br>')}</p>
    </div>
"""

        # Add findings
        html += """
    <div class="section">
        <h2>Findings</h2>
        <table>
            <tr>
                <th>Severity</th>
                <th>Vulnerability</th>
                <th>Target</th>
                <th>Description</th>
            </tr>
"""

        findings = self.results_data.get('findings', [])
        for finding in findings:
            severity_class = finding.get('severity', 'medium').lower()
            html += f"""
            <tr>
                <td class="{severity_class}">{finding.get('severity', 'N/A').upper()}</td>
                <td>{finding.get('type', 'N/A')}</td>
                <td>{finding.get('target', 'N/A')}</td>
                <td>{finding.get('description', 'N/A')}</td>
            </tr>
"""

        html += """
        </table>
    </div>
</body>
</html>
"""

        with open(path, 'w', encoding='utf-8') as f:
            f.write(html)

    def _generate_csv(self, path):
        """Generate CSV report"""
        import csv

        findings = self.results_data.get('findings', [])

        with open(path, 'w', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            writer.writerow(['Severity', 'Type', 'Target', 'Description', 'Evidence'])

            for finding in findings:
                writer.writerow([
                    finding.get('severity', ''),
                    finding.get('type', ''),
                    finding.get('target', ''),
                    finding.get('description', ''),
                    finding.get('evidence', '')
                ])

    def _generate_pdf(self, path, config):
        """Generate PDF report (requires reportlab)"""
        try:
            from reportlab.lib import colors
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
            from reportlab.lib.styles import getSampleStyleSheet

            doc = SimpleDocTemplate(str(path), pagesize=letter)
            story = []
            styles = getSampleStyleSheet()

            # Title
            title = Paragraph(config.get('report_title', 'Security Assessment Report'),
                            styles['Title'])
            story.append(title)
            story.append(Spacer(1, 20))

            # Meta info
            meta_text = f"""
            <b>Company:</b> {config.get('company_name', 'N/A')}<br/>
            <b>Client:</b> {config.get('client_name', 'N/A')}<br/>
            <b>Author:</b> {config.get('author_name', 'N/A')}<br/>
            <b>Date:</b> {datetime.now().strftime('%Y-%m-%d')}
            """
            story.append(Paragraph(meta_text, styles['Normal']))
            story.append(Spacer(1, 20))

            # Custom sections
            for section_name, section_key in [
                ('Introduction', 'intro'),
                ('Scope', 'scope'),
                ('Team', 'team'),
                ('Executive Summary', 'summary'),
                ('Methodology', 'methodology'),
                ('Process', 'process')
            ]:
                if config.get(section_key):
                    story.append(Paragraph(section_name, styles['Heading2']))
                    story.append(Paragraph(config[section_key], styles['Normal']))
                    story.append(Spacer(1, 10))

            # Findings table
            story.append(Paragraph('Findings', styles['Heading2']))

            findings = self.results_data.get('findings', [])
            if findings:
                table_data = [['Severity', 'Type', 'Target', 'Description']]
                for finding in findings:
                    table_data.append([
                        finding.get('severity', 'N/A'),
                        finding.get('type', 'N/A'),
                        finding.get('target', 'N/A')[:30],
                        finding.get('description', 'N/A')[:50]
                    ])

                t = Table(table_data)
                t.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                    ('FONTSIZE', (0, 0), (-1, -1), 8),
                ]))
                story.append(t)

            doc.build(story)
            return True

        except ImportError:
            QMessageBox.warning(
                self, "PDF Export",
                "PDF export requires 'reportlab' library.\n\n"
                "Install with: pip install reportlab"
            )
            return False

    def _generate_docx(self, path, config):
        """Generate DOCX report (requires python-docx)"""
        try:
            from docx import Document
            from docx.shared import Inches

            doc = Document()

            # Title
            doc.add_heading(config.get('report_title', 'Security Assessment Report'), 0)

            # Meta info
            doc.add_paragraph(f"Company: {config.get('company_name', 'N/A')}")
            doc.add_paragraph(f"Client: {config.get('client_name', 'N/A')}")
            doc.add_paragraph(f"Author: {config.get('author_name', 'N/A')}")
            doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d')}")

            # Custom sections
            for section_name, section_key in [
                ('Introduction', 'intro'),
                ('Scope', 'scope'),
                ('Team', 'team'),
                ('Executive Summary', 'summary'),
                ('Methodology', 'methodology'),
                ('Process', 'process')
            ]:
                if config.get(section_key):
                    doc.add_heading(section_name, 1)
                    doc.add_paragraph(config[section_key])

            # Findings
            doc.add_heading('Findings', 1)

            findings = self.results_data.get('findings', [])
            if findings:
                table = doc.add_table(rows=1, cols=4)
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Severity'
                hdr_cells[1].text = 'Type'
                hdr_cells[2].text = 'Target'
                hdr_cells[3].text = 'Description'

                for finding in findings:
                    row_cells = table.add_row().cells
                    row_cells[0].text = finding.get('severity', 'N/A')
                    row_cells[1].text = finding.get('type', 'N/A')
                    row_cells[2].text = finding.get('target', 'N/A')
                    row_cells[3].text = finding.get('description', 'N/A')

            doc.save(str(path))
            return True

        except ImportError:
            QMessageBox.warning(
                self, "DOCX Export",
                "DOCX export requires 'python-docx' library.\n\n"
                "Install with: pip install python-docx"
            )
            return False
